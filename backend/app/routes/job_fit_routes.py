"""
Job Fit & Role Matching Routes

NEW FEATURE: AI-based job fit analysis and role matching.
Analyzes resume against job descriptions to provide fit scores and recommendations.
"""

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from sqlalchemy.orm import Session
from typing import Dict, Any, List, Optional
import logging
import json
import io

from app.database import get_db
from app.ai_engines.engine_router import ai_engine_router
from app.services.resume_service import ResumeService
from app.schemas.analysis_schema import JobFitRequest, JobFitResult, RoleMatchingRequest

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/job-fit", tags=["job-fit"])

# Predefined roles for selection - Expanded list
AVAILABLE_ROLES = [
    "Software Engineer",
    "Senior Software Engineer",
    "Frontend Developer", 
    "Backend Developer",
    "Full Stack Developer",
    "Data Scientist",
    "Machine Learning Engineer",
    "AI Engineer",
    "DevOps Engineer",
    "Cloud Engineer",
    "Site Reliability Engineer",
    "Product Manager",
    "Technical Product Manager",
    "UI/UX Designer",
    "Mobile Developer",
    "iOS Developer",
    "Android Developer",
    "Cloud Architect",
    "Solutions Architect",
    "Security Engineer",
    "Cybersecurity Analyst",
    "Database Administrator",
    "Data Engineer",
    "Data Analyst",
    "Business Intelligence Analyst",
    "Quality Assurance Engineer",
    "Test Automation Engineer",
    "Technical Lead",
    "Engineering Manager",
    "CTO",
    "Blockchain Developer",
    "Game Developer",
    "Embedded Systems Engineer",
    "Network Engineer",
    "Systems Administrator",
    "Platform Engineer",
    "Research Scientist",
    "Quantitative Analyst",
    "Financial Engineer",
    "Robotics Engineer",
    "Computer Vision Engineer",
    "NLP Engineer",
    "MLOps Engineer",
    "Growth Engineer",
    "Performance Engineer",
    "Infrastructure Engineer",
    "Release Engineer",
    "Build Engineer"
]


@router.get("/available-roles")
async def get_available_roles():
    """
    Get list of available roles for job fit analysis.
    """
    return {
        "roles": AVAILABLE_ROLES,
        "total_count": len(AVAILABLE_ROLES)
    }


@router.post("/parse-resume")
async def parse_resume_for_job_fit(
    resume_file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """
    Parse resume file and extract candidate information for job fit analysis.
    
    Supports PDF, DOC, DOCX resume formats.
    Returns parsed resume data that can be used for job fit analysis.
    """
    try:
        # Validate file type
        allowed_extensions = ['.pdf', '.docx', '.doc', '.txt']
        file_extension = '.' + resume_file.filename.split('.')[-1].lower()
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Allowed types: {', '.join(allowed_extensions)}"
            )
        
        # Read file content
        resume_content = await resume_file.read()
        
        # Initialize resume service
        resume_service = ResumeService(db)
        
        # Parse resume content
        if file_extension == '.pdf':
            import pdfplumber
            with io.BytesIO(resume_content) as pdf_buffer:
                with pdfplumber.open(pdf_buffer) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() or ""
        elif file_extension in ['.docx', '.doc']:
            from docx import Document
            with io.BytesIO(resume_content) as docx_buffer:
                doc = Document(docx_buffer)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        else:  # .txt
            text = resume_content.decode('utf-8')
        
        # Extract structured data from text
        parsed_data = resume_service._extract_resume_data(text)
        
        # Validate parsed data
        missing_fields = resume_service.validate_resume(parsed_data)
        
        return {
            "success": True,
            "filename": resume_file.filename,
            "parsed_data": parsed_data,
            "validation": {
                "is_valid": len(missing_fields) == 0,
                "missing_fields": missing_fields
            },
            "available_roles": AVAILABLE_ROLES,
            "estimated_role": parsed_data.get("estimated_role", "Software Engineer")
        }
        
    except Exception as e:
        logger.error(f"Error parsing resume: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to parse resume: {str(e)}")


@router.post("/analyze-with-role")
async def analyze_job_fit_with_role(
    parsed_resume: str = Form(...),
    selected_role: str = Form(...)
):
    """
    Analyze job fit between parsed resume data and selected role using Ollama AI.
    
    This endpoint takes parsed resume data and a selected role, then uses Ollama
    to determine if the candidate is a good fit for the role.
    """
    try:
        # Validate selected role (only for predefined roles)
        if selected_role in AVAILABLE_ROLES:
            # It's a predefined role, proceed normally
            pass
        else:
            # It's a custom role, we'll create a generic job description
            logger.info(f"Using custom role: {selected_role}")
        
        # Parse the resume data from JSON string
        try:
            parsed_resume_data = json.loads(parsed_resume)
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="Invalid resume data format")
        
        # Create a job description based on the selected role
        job_description = _create_role_based_job_description(selected_role)
        
        # Use Ollama via AI Engine Router for job fit analysis
        logger.info(f"Analyzing job fit for {selected_role} using Ollama")
        
        fit_analysis = ai_engine_router.calculate_job_fit(
            candidate_context=parsed_resume_data,
            job_description=job_description
        )
        
        # Enhance the analysis with role-specific insights
        enhanced_analysis = _enhance_job_fit_analysis(
            fit_analysis, 
            parsed_resume_data, 
            selected_role
        )
        
        return {
            "success": True,
            "role": selected_role,
            "candidate_summary": {
                "name": "Candidate",  # Could extract from resume if available
                "experience_years": parsed_resume_data.get("experience_years", 0),
                "experience_level": parsed_resume_data.get("experience", {}).get("level", "Junior"),
                "top_skills": parsed_resume_data.get("skills", [])[:5],
                "estimated_role": parsed_resume_data.get("estimated_role", "Software Engineer")
            },
            "job_fit_analysis": enhanced_analysis,
            "recommendation": _generate_job_fit_recommendation(enhanced_analysis),
            "next_steps": _generate_next_steps(enhanced_analysis, selected_role)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing job fit: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to analyze job fit: {str(e)}")


@router.post("/bulk-role-analysis")
async def analyze_multiple_roles(
    parsed_resume: str = Form(...),
    roles: str = Form(...)
):
    """
    Analyze job fit for multiple roles at once.
    Returns ranked list of roles based on fit scores.
    """
    try:
        # Parse the resume data from JSON string
        try:
            parsed_resume_data = json.loads(parsed_resume)
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="Invalid resume data format")
        
        # Parse roles from JSON string or comma-separated string
        try:
            if roles.startswith('['):
                roles_list = json.loads(roles)
            else:
                roles_list = [role.strip() for role in roles.split(',')]
        except (json.JSONDecodeError, AttributeError):
            raise HTTPException(status_code=400, detail="Invalid roles format")
        
        # Validate roles
        invalid_roles = [role for role in roles_list if role not in AVAILABLE_ROLES]
        if invalid_roles:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid roles: {', '.join(invalid_roles)}"
            )
        
        role_analyses = []
        
        for role in roles_list:
            try:
                # Create job description for this role
                job_description = _create_role_based_job_description(role)
                
                # Analyze fit using Ollama
                fit_analysis = ai_engine_router.calculate_job_fit(
                    candidate_context=parsed_resume_data,
                    job_description=job_description
                )
                
                role_analyses.append({
                    "role": role,
                    "overall_fit_score": fit_analysis.get("overall_fit_score", 0),
                    "skill_match_percentage": fit_analysis.get("skill_match_percentage", 0),
                    "experience_match_percentage": fit_analysis.get("experience_match_percentage", 0),
                    "role_suitability": fit_analysis.get("role_suitability", "Unknown"),
                    "top_missing_skills": fit_analysis.get("missing_skills", [])[:3],
                    "top_matched_skills": fit_analysis.get("matched_skills", [])[:5]
                })
                
            except Exception as e:
                logger.error(f"Error analyzing role {role}: {e}")
                role_analyses.append({
                    "role": role,
                    "error": f"Analysis failed: {str(e)}",
                    "overall_fit_score": 0
                })
        
        # Sort by fit score (descending)
        role_analyses.sort(key=lambda x: x.get("overall_fit_score", 0), reverse=True)
        
        return {
            "success": True,
            "candidate_summary": {
                "experience_years": parsed_resume_data.get("experience_years", 0),
                "experience_level": parsed_resume_data.get("experience", {}).get("level", "Junior"),
                "top_skills": parsed_resume_data.get("skills", [])[:5]
            },
            "role_analyses": role_analyses,
            "best_fit_role": role_analyses[0]["role"] if role_analyses else None,
            "total_roles_analyzed": len(role_analyses)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in bulk role analysis: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to analyze multiple roles: {str(e)}")


def _create_role_based_job_description(role: str) -> Dict[str, Any]:
    """
    Create a comprehensive job description based on the selected role.
    """
    role_descriptions = {
        "Software Engineer": {
            "title": "Software Engineer",
            "required_skills": ["Programming", "Problem Solving", "Software Development", "Version Control", "Testing", "Debugging"],
            "preferred_skills": ["Python", "Java", "JavaScript", "SQL", "Git", "Agile", "CI/CD", "REST API", "Docker"],
            "required_experience_years": 2,
            "description": "Develop, test, and maintain software applications using modern programming languages and frameworks."
        },
        "Senior Software Engineer": {
            "title": "Senior Software Engineer",
            "required_skills": ["Advanced Programming", "System Design", "Architecture", "Code Review", "Mentoring", "Technical Leadership"],
            "preferred_skills": ["Python", "Java", "JavaScript", "Microservices", "Cloud Platforms", "Database Design", "Performance Optimization"],
            "required_experience_years": 5,
            "description": "Lead software development projects, mentor junior developers, and design scalable systems."
        },
        "Frontend Developer": {
            "title": "Frontend Developer", 
            "required_skills": ["HTML", "CSS", "JavaScript", "Responsive Design", "UI/UX", "Cross-browser Compatibility"],
            "preferred_skills": ["React", "Vue.js", "Angular", "TypeScript", "SASS", "Webpack", "Git", "Figma", "Bootstrap", "Tailwind CSS"],
            "required_experience_years": 2,
            "description": "Build user-facing web applications with modern frontend technologies and frameworks."
        },
        "Backend Developer": {
            "title": "Backend Developer",
            "required_skills": ["Server-side Programming", "Database Design", "API Development", "System Architecture", "Security"],
            "preferred_skills": ["Python", "Java", "Node.js", "SQL", "NoSQL", "REST API", "Microservices", "Docker", "AWS", "Redis"],
            "required_experience_years": 3,
            "description": "Design and implement server-side logic, databases, and APIs for web applications."
        },
        "Full Stack Developer": {
            "title": "Full Stack Developer",
            "required_skills": ["Frontend Development", "Backend Development", "Database Management", "System Design", "API Integration"],
            "preferred_skills": ["JavaScript", "Python", "React", "Node.js", "SQL", "MongoDB", "Git", "Docker", "AWS", "TypeScript"],
            "required_experience_years": 3,
            "description": "Work on both frontend and backend components of web applications."
        },
        "Data Scientist": {
            "title": "Data Scientist",
            "required_skills": ["Statistics", "Machine Learning", "Data Analysis", "Programming", "Data Visualization", "Research"],
            "preferred_skills": ["Python", "R", "SQL", "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "Tableau", "Jupyter", "Statistics"],
            "required_experience_years": 3,
            "description": "Analyze complex data sets to derive insights and build predictive models."
        },
        "Machine Learning Engineer": {
            "title": "Machine Learning Engineer",
            "required_skills": ["Machine Learning", "Deep Learning", "Model Deployment", "Programming", "Statistics", "MLOps"],
            "preferred_skills": ["Python", "TensorFlow", "PyTorch", "Scikit-learn", "Docker", "Kubernetes", "MLOps", "AWS", "Feature Engineering"],
            "required_experience_years": 4,
            "description": "Design, build, and deploy machine learning models and systems at scale."
        },
        "AI Engineer": {
            "title": "AI Engineer",
            "required_skills": ["Artificial Intelligence", "Machine Learning", "Deep Learning", "Neural Networks", "Algorithm Design"],
            "preferred_skills": ["Python", "TensorFlow", "PyTorch", "Computer Vision", "NLP", "Reinforcement Learning", "MLOps", "Cloud AI Services"],
            "required_experience_years": 4,
            "description": "Develop and implement AI solutions and intelligent systems."
        },
        "DevOps Engineer": {
            "title": "DevOps Engineer",
            "required_skills": ["CI/CD", "Infrastructure Management", "Automation", "Monitoring", "Cloud Platforms", "Containerization"],
            "preferred_skills": ["AWS", "Docker", "Kubernetes", "Jenkins", "Terraform", "Ansible", "Linux", "Git", "Prometheus", "Grafana"],
            "required_experience_years": 3,
            "description": "Manage infrastructure, automate deployments, and ensure system reliability."
        },
        "Cloud Engineer": {
            "title": "Cloud Engineer",
            "required_skills": ["Cloud Platforms", "Infrastructure as Code", "Networking", "Security", "Automation", "Monitoring"],
            "preferred_skills": ["AWS", "Azure", "GCP", "Terraform", "CloudFormation", "Docker", "Kubernetes", "Linux", "Python", "Bash"],
            "required_experience_years": 3,
            "description": "Design, implement, and manage cloud infrastructure and services."
        },
        "Site Reliability Engineer": {
            "title": "Site Reliability Engineer",
            "required_skills": ["System Reliability", "Monitoring", "Incident Response", "Automation", "Performance Optimization"],
            "preferred_skills": ["Linux", "Python", "Go", "Kubernetes", "Prometheus", "Grafana", "Terraform", "AWS", "Incident Management"],
            "required_experience_years": 4,
            "description": "Ensure system reliability, performance, and availability at scale."
        },
        "Product Manager": {
            "title": "Product Manager",
            "required_skills": ["Product Strategy", "Market Research", "Project Management", "Communication", "Analytics", "User Research"],
            "preferred_skills": ["Agile", "Scrum", "User Research", "A/B Testing", "SQL", "Data Analysis", "Roadmapping", "Stakeholder Management"],
            "required_experience_years": 4,
            "description": "Define product strategy, manage roadmaps, and coordinate cross-functional teams."
        },
        "Technical Product Manager": {
            "title": "Technical Product Manager",
            "required_skills": ["Product Management", "Technical Knowledge", "API Design", "System Architecture", "Engineering Collaboration"],
            "preferred_skills": ["Programming", "API Design", "Microservices", "Cloud Platforms", "Data Analysis", "Technical Writing"],
            "required_experience_years": 5,
            "description": "Bridge technical and business requirements for complex technical products."
        },
        "UI/UX Designer": {
            "title": "UI/UX Designer",
            "required_skills": ["User Experience Design", "User Interface Design", "Prototyping", "User Research", "Design Thinking"],
            "preferred_skills": ["Figma", "Sketch", "Adobe Creative Suite", "Wireframing", "Usability Testing", "Design Systems", "HTML/CSS"],
            "required_experience_years": 2,
            "description": "Design intuitive and engaging user interfaces and experiences."
        },
        "Mobile Developer": {
            "title": "Mobile Developer",
            "required_skills": ["Mobile Development", "UI Design", "API Integration", "App Store Guidelines", "Performance Optimization"],
            "preferred_skills": ["React Native", "Flutter", "Swift", "Kotlin", "iOS", "Android", "Firebase", "Push Notifications"],
            "required_experience_years": 3,
            "description": "Develop mobile applications for iOS and Android platforms."
        },
        "iOS Developer": {
            "title": "iOS Developer",
            "required_skills": ["iOS Development", "Swift", "Objective-C", "Xcode", "App Store Guidelines", "iOS SDK"],
            "preferred_skills": ["SwiftUI", "UIKit", "Core Data", "Firebase", "Push Notifications", "In-App Purchases", "TestFlight"],
            "required_experience_years": 3,
            "description": "Develop native iOS applications using Swift and iOS frameworks."
        },
        "Android Developer": {
            "title": "Android Developer",
            "required_skills": ["Android Development", "Kotlin", "Java", "Android Studio", "Google Play Guidelines", "Android SDK"],
            "preferred_skills": ["Jetpack Compose", "Room Database", "Firebase", "Material Design", "Retrofit", "Dagger/Hilt"],
            "required_experience_years": 3,
            "description": "Develop native Android applications using Kotlin and Android frameworks."
        },
        "Data Engineer": {
            "title": "Data Engineer",
            "required_skills": ["Data Pipeline", "ETL", "Big Data", "Database Design", "Data Warehousing", "Programming"],
            "preferred_skills": ["Python", "SQL", "Spark", "Kafka", "Airflow", "AWS", "Snowflake", "dbt", "Hadoop", "Elasticsearch"],
            "required_experience_years": 3,
            "description": "Build and maintain data infrastructure and pipelines for analytics and ML."
        },
        "Data Analyst": {
            "title": "Data Analyst",
            "required_skills": ["Data Analysis", "SQL", "Statistics", "Data Visualization", "Business Intelligence", "Reporting"],
            "preferred_skills": ["Python", "R", "Tableau", "Power BI", "Excel", "Pandas", "Statistical Analysis", "A/B Testing"],
            "required_experience_years": 2,
            "description": "Analyze data to provide business insights and support decision-making."
        },
        "Security Engineer": {
            "title": "Security Engineer",
            "required_skills": ["Cybersecurity", "Network Security", "Vulnerability Assessment", "Incident Response", "Security Frameworks"],
            "preferred_skills": ["Penetration Testing", "OWASP", "SIEM", "Firewall", "Encryption", "Compliance", "Risk Assessment"],
            "required_experience_years": 4,
            "description": "Implement and maintain security measures to protect systems and data."
        },
        "Blockchain Developer": {
            "title": "Blockchain Developer",
            "required_skills": ["Blockchain Technology", "Smart Contracts", "Cryptocurrency", "Distributed Systems", "Cryptography"],
            "preferred_skills": ["Solidity", "Ethereum", "Web3", "DeFi", "NFT", "Hyperledger", "Truffle", "Ganache", "JavaScript"],
            "required_experience_years": 3,
            "description": "Develop blockchain applications and smart contracts."
        },
        "Game Developer": {
            "title": "Game Developer",
            "required_skills": ["Game Development", "Game Design", "Programming", "Game Engines", "Graphics Programming"],
            "preferred_skills": ["Unity", "Unreal Engine", "C#", "C++", "Game Physics", "3D Graphics", "Mobile Games", "VR/AR"],
            "required_experience_years": 3,
            "description": "Design and develop video games for various platforms."
        },
        "MLOps Engineer": {
            "title": "MLOps Engineer",
            "required_skills": ["MLOps", "Machine Learning", "DevOps", "Model Deployment", "CI/CD for ML", "Monitoring"],
            "preferred_skills": ["Python", "Docker", "Kubernetes", "MLflow", "Kubeflow", "AWS SageMaker", "TensorFlow Serving", "Monitoring"],
            "required_experience_years": 4,
            "description": "Operationalize machine learning models and manage ML infrastructure."
        },
        "Computer Vision Engineer": {
            "title": "Computer Vision Engineer",
            "required_skills": ["Computer Vision", "Image Processing", "Deep Learning", "OpenCV", "Neural Networks"],
            "preferred_skills": ["Python", "TensorFlow", "PyTorch", "OpenCV", "YOLO", "CNN", "Object Detection", "Image Segmentation"],
            "required_experience_years": 4,
            "description": "Develop computer vision systems and image processing applications."
        },
        "NLP Engineer": {
            "title": "NLP Engineer",
            "required_skills": ["Natural Language Processing", "Machine Learning", "Text Processing", "Linguistics", "Deep Learning"],
            "preferred_skills": ["Python", "NLTK", "spaCy", "Transformers", "BERT", "GPT", "Hugging Face", "TensorFlow", "PyTorch"],
            "required_experience_years": 4,
            "description": "Build natural language processing systems and text analysis applications."
        }
    }
    
    # Return the job description for the role, or a generic one if not found
    return role_descriptions.get(role, {
        "title": role,
        "required_skills": ["Programming", "Problem Solving", "Communication", "Technical Skills"],
        "preferred_skills": ["Relevant Technical Skills", "Team Collaboration", "Continuous Learning"],
        "required_experience_years": 2,
        "description": f"Professional role in {role} with relevant technical skills and experience."
    })


def _enhance_job_fit_analysis(fit_analysis: Dict[str, Any], resume_data: Dict[str, Any], role: str) -> Dict[str, Any]:
    """
    Enhance the basic job fit analysis with additional insights.
    """
    enhanced = fit_analysis.copy()
    
    # Add role-specific analysis
    enhanced["role_specific_insights"] = {
        "experience_alignment": _analyze_experience_alignment(resume_data, role),
        "skill_gap_analysis": _analyze_skill_gaps(resume_data, role),
        "growth_potential": _analyze_growth_potential(resume_data, role),
        "cultural_fit_indicators": _analyze_cultural_fit(resume_data, role)
    }
    
    # Add confidence score
    enhanced["confidence_score"] = _calculate_confidence_score(fit_analysis)
    
    return enhanced


def _analyze_experience_alignment(resume_data: Dict[str, Any], role: str) -> Dict[str, Any]:
    """Analyze how well the candidate's experience aligns with the role."""
    experience_years = resume_data.get("experience_years", 0)
    estimated_role = resume_data.get("estimated_role", "")
    
    role_similarity = 0.8 if role.lower() in estimated_role.lower() else 0.3
    
    return {
        "years_experience": experience_years,
        "role_similarity": role_similarity,
        "alignment_score": min(100, (experience_years * 20) + (role_similarity * 50))
    }


def _analyze_skill_gaps(resume_data: Dict[str, Any], role: str) -> Dict[str, Any]:
    """Analyze skill gaps and learning opportunities."""
    candidate_skills = [skill.lower() for skill in resume_data.get("skills", [])]
    job_desc = _create_role_based_job_description(role)
    required_skills = [skill.lower() for skill in job_desc.get("required_skills", [])]
    
    missing_critical = [skill for skill in required_skills if skill not in candidate_skills]
    
    return {
        "critical_gaps": missing_critical[:3],
        "learning_priority": "High" if len(missing_critical) > 2 else "Medium" if missing_critical else "Low",
        "estimated_learning_time": f"{len(missing_critical) * 2}-{len(missing_critical) * 4} months"
    }


def _analyze_growth_potential(resume_data: Dict[str, Any], role: str) -> Dict[str, Any]:
    """Analyze the candidate's growth potential in the role."""
    experience_years = resume_data.get("experience_years", 0)
    skills_count = len(resume_data.get("skills", []))
    projects_count = len(resume_data.get("projects", []))
    
    growth_score = min(100, (skills_count * 3) + (projects_count * 5) + (experience_years * 10))
    
    return {
        "growth_score": growth_score,
        "potential_level": "High" if growth_score > 70 else "Medium" if growth_score > 40 else "Low",
        "development_areas": ["Technical Skills", "Leadership", "Domain Expertise"][:2]
    }


def _analyze_cultural_fit(resume_data: Dict[str, Any], role: str) -> Dict[str, Any]:
    """Analyze cultural fit indicators."""
    projects = resume_data.get("projects", [])
    
    collaboration_indicators = sum(1 for project in projects if any(
        word in project.lower() for word in ["team", "collaboration", "group", "pair"]
    ))
    
    return {
        "collaboration_score": min(100, collaboration_indicators * 25),
        "leadership_indicators": collaboration_indicators > 2,
        "adaptability_score": 75  # Default score, could be enhanced with more analysis
    }


def _calculate_confidence_score(fit_analysis: Dict[str, Any]) -> int:
    """Calculate confidence score for the analysis."""
    overall_score = fit_analysis.get("overall_fit_score", 0)
    skill_match = fit_analysis.get("skill_match_percentage", 0)
    experience_match = fit_analysis.get("experience_match_percentage", 0)
    
    # Higher confidence when scores are consistent
    score_variance = abs(overall_score - skill_match) + abs(overall_score - experience_match)
    confidence = max(50, 100 - (score_variance / 2))
    
    return int(confidence)


def _generate_job_fit_recommendation(analysis: Dict[str, Any]) -> Dict[str, Any]:
    """Generate a recommendation based on the job fit analysis."""
    overall_score = analysis.get("overall_fit_score", 0)
    confidence = analysis.get("confidence_score", 50)
    
    if overall_score >= 80:
        recommendation = "Excellent Fit"
        action = "Strongly recommend for this role"
        color = "green"
    elif overall_score >= 65:
        recommendation = "Good Fit"
        action = "Recommend with minor skill development"
        color = "blue"
    elif overall_score >= 50:
        recommendation = "Moderate Fit"
        action = "Consider with additional training"
        color = "yellow"
    else:
        recommendation = "Poor Fit"
        action = "Not recommended for this role"
        color = "red"
    
    return {
        "recommendation": recommendation,
        "action": action,
        "confidence_level": "High" if confidence > 75 else "Medium" if confidence > 50 else "Low",
        "color": color,
        "score": overall_score
    }


def _generate_next_steps(analysis: Dict[str, Any], role: str) -> List[str]:
    """Generate actionable next steps based on the analysis."""
    steps = []
    overall_score = analysis.get("overall_fit_score", 0)
    missing_skills = analysis.get("missing_skills", [])
    
    if overall_score >= 80:
        steps.extend([
            "Proceed with technical interview",
            "Discuss role expectations and responsibilities",
            "Review compensation and benefits"
        ])
    elif overall_score >= 50:
        if missing_skills:
            steps.append(f"Develop skills in: {', '.join(missing_skills[:3])}")
        steps.extend([
            "Consider additional training or certification",
            "Schedule follow-up interview after skill development",
            "Explore mentorship opportunities"
        ])
    else:
        steps.extend([
            "Focus on fundamental skill development",
            "Consider alternative roles that better match current skills",
            "Gain more experience through projects or internships"
        ])
    
    return steps


@router.post("/upload-resume-analyze")
async def analyze_job_fit_from_resume(
    resume_file: UploadFile = File(...),
    job_description: str = None,
    db: Session = Depends(get_db)
):
    """
    Upload resume file and analyze job fit against provided job description.
    
    Supports PDF, DOC, DOCX resume formats.
    """
    try:
        if not job_description:
            raise HTTPException(status_code=400, detail="Job description is required")
        
        # Parse job description JSON
        try:
            job_desc_data = json.loads(job_description)
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="Invalid job description JSON format")
        
        # Read and parse resume
        resume_content = await resume_file.read()
        resume_service = ResumeService(db)
        
        # Parse resume content using the same method as parse_resume_for_job_fit
        if resume_file.filename.endswith('.pdf'):
            import pdfplumber
            with io.BytesIO(resume_content) as pdf_buffer:
                with pdfplumber.open(pdf_buffer) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() or ""
        elif resume_file.filename.endswith(('.docx', '.doc')):
            from docx import Document
            with io.BytesIO(resume_content) as docx_buffer:
                doc = Document(docx_buffer)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        else:
            text = resume_content.decode('utf-8')
        
        parsed_resume = resume_service._extract_resume_data(text)
        
        # Analyze job fit using AI Engine Router (Ollama)
        fit_analysis = ai_engine_router.calculate_job_fit(
            candidate_context=parsed_resume,
            job_description=job_desc_data
        )
        
        return {
            "resume_analysis": parsed_resume,
            "job_fit_analysis": fit_analysis,
            "filename": resume_file.filename
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing resume job fit: {e}")
        raise HTTPException(status_code=500, detail="Failed to analyze resume job fit")


@router.post("/role-matching", response_model=List[Dict[str, Any]])
async def find_matching_roles(
    request: RoleMatchingRequest
):
    """
    Find best matching roles for a candidate based on their profile.
    
    Returns ranked list of roles with fit scores and explanations.
    """
    try:
        logger.info("Finding matching roles for candidate profile")
        
        # Use predefined roles instead of sample database
        role_matches = []
        for role in AVAILABLE_ROLES[:8]:  # Analyze top 8 roles
            try:
                # Create job description for this role
                job_description = _create_role_based_job_description(role)
                
                # Analyze fit using AI Engine Router (Ollama)
                fit_analysis = ai_engine_router.calculate_job_fit(
                    candidate_context=request.candidate_profile,
                    job_description=job_description
                )
                
                role_matches.append({
                    "role": {
                        "title": role,
                        "required_skills": job_description["required_skills"],
                        "required_experience_years": job_description["required_experience_years"]
                    },
                    "fit_score": fit_analysis.get('overall_fit_score', 0),
                    "skill_match": fit_analysis.get('skill_match_percentage', 0),
                    "experience_match": fit_analysis.get('experience_match_percentage', 0),
                    "suitability": fit_analysis.get('role_suitability', 'Unknown'),
                    "missing_skills": fit_analysis.get('missing_skills', [])[:3],
                    "recommendations": ["Focus on skill development", "Gain relevant experience"][:2]
                })
                
            except Exception as e:
                logger.error(f"Error analyzing role {role}: {e}")
                continue
        
        # Sort by fit score (descending)
        role_matches.sort(key=lambda x: x['fit_score'], reverse=True)
        
        return role_matches
        
    except Exception as e:
        logger.error(f"Error finding matching roles: {e}")
        raise HTTPException(status_code=500, detail="Failed to find matching roles")


@router.get("/sample-job-descriptions")
async def get_sample_job_descriptions():
    """
    Get sample job descriptions for testing job fit analysis.
    """
    return {
        "sample_jobs": [
            {
                "title": "Senior Software Engineer",
                "company": "TechCorp",
                "required_skills": ["Python", "JavaScript", "React", "Node.js", "PostgreSQL"],
                "preferred_skills": ["AWS", "Docker", "TypeScript", "GraphQL"],
                "required_experience_years": 5,
                "education_requirements": {
                    "degree": "Bachelor's",
                    "field": "Computer Science or related"
                },
                "description": "We're looking for a senior software engineer to join our growing team..."
            },
            {
                "title": "Frontend Developer",
                "company": "StartupXYZ", 
                "required_skills": ["JavaScript", "React", "HTML", "CSS"],
                "preferred_skills": ["TypeScript", "Next.js", "Tailwind CSS"],
                "required_experience_years": 3,
                "education_requirements": {
                    "degree": "Bachelor's or equivalent experience"
                },
                "description": "Join our frontend team to build amazing user experiences..."
            },
            {
                "title": "Data Scientist",
                "company": "DataCorp",
                "required_skills": ["Python", "Machine Learning", "SQL", "Statistics"],
                "preferred_skills": ["TensorFlow", "PyTorch", "AWS", "Spark"],
                "required_experience_years": 4,
                "education_requirements": {
                    "degree": "Master's preferred",
                    "field": "Data Science, Statistics, or related"
                },
                "description": "We're seeking a data scientist to drive insights from our data..."
            }
        ]
    }


# Helper functions for detailed analysis

def _analyze_skill_categories(candidate_skills: List[str], required_skills: List[str], preferred_skills: List[str]) -> Dict[str, Any]:
    """Analyze skills by category (programming languages, frameworks, tools, etc.)"""
    
    categories = {
        "programming_languages": ["Python", "JavaScript", "Java", "C++", "Go", "Rust", "TypeScript"],
        "frameworks": ["React", "Angular", "Vue.js", "Django", "Flask", "Express", "Spring"],
        "databases": ["PostgreSQL", "MySQL", "MongoDB", "Redis", "Elasticsearch"],
        "cloud_platforms": ["AWS", "Azure", "GCP", "Docker", "Kubernetes"],
        "tools": ["Git", "Jenkins", "Terraform", "Ansible", "Webpack"]
    }
    
    candidate_lower = [skill.lower() for skill in candidate_skills]
    required_lower = [skill.lower() for skill in required_skills]
    preferred_lower = [skill.lower() for skill in preferred_skills]
    
    analysis = {}
    
    for category, skills in categories.items():
        skills_lower = [skill.lower() for skill in skills]
        
        candidate_in_category = [skill for skill in candidate_lower if skill in skills_lower]
        required_in_category = [skill for skill in required_lower if skill in skills_lower]
        preferred_in_category = [skill for skill in preferred_lower if skill in skills_lower]
        
        matches = set(candidate_in_category).intersection(set(required_in_category + preferred_in_category))
        
        analysis[category] = {
            "candidate_skills": candidate_in_category,
            "required_matches": list(set(candidate_in_category).intersection(set(required_in_category))),
            "preferred_matches": list(set(candidate_in_category).intersection(set(preferred_in_category))),
            "total_matches": len(matches)
        }
    
    return analysis


def _analyze_experience_fit(candidate_years: int, required_years: int, projects: List[Dict]) -> Dict[str, Any]:
    """Analyze experience fit including years and project relevance"""
    
    experience_gap = required_years - candidate_years
    experience_ratio = candidate_years / max(required_years, 1)
    
    # Analyze project complexity and relevance
    project_score = 0
    if projects:
        # Simple scoring based on project descriptions
        for project in projects[:5]:  # Top 5 projects
            description = project.get('description', '').lower()
            if any(word in description for word in ['led', 'managed', 'architected', 'designed']):
                project_score += 2
            if any(word in description for word in ['team', 'collaboration', 'cross-functional']):
                project_score += 1
            if any(word in description for word in ['scale', 'performance', 'optimization']):
                project_score += 1
    
    return {
        "years_experience": candidate_years,
        "required_years": required_years,
        "experience_gap": experience_gap,
        "experience_ratio": round(experience_ratio, 2),
        "project_leadership_score": project_score,
        "assessment": "Exceeds requirements" if experience_gap <= -2 else
                    "Meets requirements" if experience_gap <= 0 else
                    "Below requirements" if experience_gap <= 2 else
                    "Significantly below requirements"
    }


def _analyze_education_fit(candidate_education: Dict, required_education: Dict) -> Dict[str, Any]:
    """Analyze education requirements fit"""
    
    if not required_education:
        return {"match": True, "notes": "No specific education requirements"}
    
    candidate_degree = candidate_education.get('degree', '').lower()
    required_degree = required_education.get('degree', '').lower()
    
    degree_hierarchy = {
        'phd': 4, 'doctorate': 4,
        'master': 3, 'masters': 3, 'msc': 3, 'mba': 3,
        'bachelor': 2, 'bachelors': 2, 'bsc': 2, 'ba': 2,
        'associate': 1, 'diploma': 1
    }
    
    candidate_level = 0
    required_level = 0
    
    for degree, level in degree_hierarchy.items():
        if degree in candidate_degree:
            candidate_level = max(candidate_level, level)
        if degree in required_degree:
            required_level = max(required_level, level)
    
    meets_requirement = candidate_level >= required_level
    
    return {
        "candidate_degree": candidate_education.get('degree', 'Not specified'),
        "required_degree": required_education.get('degree', 'Not specified'),
        "meets_requirement": meets_requirement,
        "education_level_match": candidate_level >= required_level,
        "field_match": _check_field_match(
            candidate_education.get('field', ''),
            required_education.get('field', '')
        )
    }


def _check_field_match(candidate_field: str, required_field: str) -> bool:
    """Check if education field matches requirements"""
    if not required_field or not candidate_field:
        return True  # No specific requirement
    
    candidate_lower = candidate_field.lower()
    required_lower = required_field.lower()
    
    # Simple keyword matching - could be enhanced with more sophisticated matching
    related_fields = {
        'computer science': ['software', 'programming', 'computing', 'informatics'],
        'engineering': ['technical', 'technology', 'science'],
        'business': ['management', 'administration', 'commerce'],
        'data science': ['statistics', 'mathematics', 'analytics', 'data']
    }
    
    # Direct match
    if any(word in candidate_lower for word in required_lower.split()):
        return True
    
    # Related field match
    for field, related in related_fields.items():
        if field in required_lower and any(word in candidate_lower for word in related):
            return True
    
    return False